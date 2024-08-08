#!/usr/bin/env python
import sys                  # To get terminal input
import shutil
import requests, bs4        # To fetch data from the online profiles
import docx                 # To make the .docx file

if len(sys.argv) > 1:
    address = sys.argv[1]

res = requests.get(address)
res.raise_for_status()
helper = bs4.BeautifulSoup(res.text, 'html.parser')

def get_name(soup):
    name = soup.select(".mb-0")[0].text.strip()
    return name
def get_age(soup):
    name_age = soup.select("div.user-detail")[0].text
    age_w_stuff = name_age.split("(")[1]
    age = age_w_stuff.split(" ")[0]
    return age
def get_marital(soup):
    line = soup.select("div.hp-candidate-wrapper")[0].text
    line = line.split("|")
    marital = line[1].strip()
    return marital
def get_kids(soup):
    line = soup.select("div.hp-candidate-wrapper")[0].text
    line = line.split("|")
    kids = line[2].strip()
    return kids
def get_religion(soup):
    line = soup.select("div.hp-candidate-wrapper")[0].text
    line = line.split("|")
    religion_contact = line[-1].strip()
    religion = religion_contact.split(' ')[0]
    return religion
def get_reason(soup):
    block = soup.select("h3.footer-experience")[0].text
    reason = block.split("|")[1].strip()
    return reason
def get_years_exp(soup):
    block = soup.select("h3.footer-experience")[2].text
    yrs_exp = block.strip().split(" ")[0]
    return yrs_exp
def get_visa_end(soup):
    block = soup.select("h3.footer-experience")[3].text
    block = block.split("|")[0]
    day = block.split(" ")[3]
    month = block.split(" ")[4]
    return day + " " + month

def get_main_skills(soup):
    row = soup.select("h4.float-left.color_2")
    skills = [row[0].text.strip()] + [ item.text.strip().lower() for item in row[1:] ]
    return ", ".join(skills)
def get_cook_skills(soup):
    row = soup.select("h4.float-left.color_3")
    skills = [row[0].text.strip()] + [ item.text.strip().lower() for item in row[1:] ]
    return ", ".join(skills)
def get_other_skills(soup):
    row = soup.select("h4.float-left.color_4")
    skills = [row[0].text.strip()] + [ item.text.strip().lower() for item in row[1:] ]
    return ", ".join(skills)
def get_personality(soup):
    row = soup.select("h4.float-left.color_5")
    personality = [row[0].text.strip()] + [ item.text.strip().lower() for item in row[1:] ]
    return ", ".join(personality)

def get_about_me(soup):
    aboutme = soup.select("p")[1].text
    return aboutme

def get_work_experience(soup):
    months = ['Jan','Feb','Mar','Apr','May','Jun',
              'Jul','Aug','Sep','Oct','Nov','Dec']
    exps = ''
    current_exp = ''
    ind = 1
    try:
        leading_word = soup.select(".mb-0")[ind].text.split(' ')[1]
    except:
        leading_word = None

    while leading_word in months:
        current_exp += soup.select(".mb-0")[ind].text.strip()
        current_exp += '\n'
        current_exp += soup.select(".mb-0")[ind+1].text.strip()
        current_exp += '\n'
        current_exp += soup.select(".mb-1")[3*((ind-1)//2)].text.strip()
        current_exp += '\n'
        current_exp += soup.select(".mb-1")[3*((ind-1)//2)+1].text.strip()
        current_exp += '\n'
        current_exp += soup.select(".mb-1")[3*((ind-1)//2)+2].text.strip()

        if len(exps) > 0:
            exps += '\n\n'
        exps += current_exp
        current_exp = ''

        ind += 2
        try:
            leading_word = soup.select(".mb-0")[ind].text.split(' ')[1]
        except:
            leading_word = None


    return exps


title = get_name(helper).split(" ")[0].lower()
shutil.copy("template.docx", f"{title}.docx")
doc = docx.Document(f"{title}.docx")
for para in doc.paragraphs:
    if 'NAME' in para.text:
        para.text = ''
        para.add_run('NAME:')
        para.runs[0].underline = True
        para.runs[0].bold = True
        para.add_run(f' {get_name(helper)}')
        para.style = 'Heading 1'
    if 'AGE' in para.text:
        para.text = f"Age: {get_age(helper)} y/o"
        para.style = 'Normal'
    if 'MARITAL-STATUS' in para.text:
        para.text = f"Marital status: {get_marital(helper)} with {get_kids(helper)}"
        para.style = 'Normal'
    if 'RELIGION' in para.text:
        para.text = f"Religion: {get_religion(helper)}"
        para.style = 'Normal'
    if 'REASON-OF-LEAVING' in para.text:
        para.text = f"Reason of leaving: {get_reason(helper)}"
        para.style = 'Normal'
    if 'YEARS-EXPERIENCE' in para.text:
        para.text = f"{get_years_exp(helper)} years work experience"
        para.style = 'Normal'
    if 'VISA-LAST-DAY' in para.text:
        para.text = f"Last day of visa: {get_visa_end(helper)}"
        para.style = 'Normal'
    if 'MAIN-SKILLS' in para.text:
        para.text = f"Main Skills: {get_main_skills(helper)}"
        para.style = 'Normal'
    if 'COOKING-SKILLS' in para.text:
        para.text = f"Cooking Skills: {get_cook_skills(helper)}"
        para.style = 'Normal'
    if 'OTHER-SKILLS' in para.text:
        para.text = f"Other Skills: {get_other_skills(helper)}"
        para.style = 'Normal'
    if 'PERSONALITY' in para.text:
        para.text = f"Personality: {get_personality(helper)}"
        para.style = 'Normal'
    if 'ABOUT-ME' in para.text:
        para.text = get_about_me(helper)
        para.style = 'Normal'
    if 'WORK-EXPERIENCE' in para.text:
        para.text = get_work_experience(helper)
        para.style = 'Normal'

doc.save(f"{title}.docx")
